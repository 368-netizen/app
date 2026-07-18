import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import io

def create_uew_document(data):
    doc = docx.Document()
    
    # Set UEW Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Configure UEW standard font and spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0  # Double spacing
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def add_centered_line(text, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold

    # --- Cover Page Layout ---
    add_centered_line("UNIVERSITY OF EDUCATION, WINNEBA", bold=True)
    add_centered_line(data['faculty'].upper())
    add_centered_line(data['department'].upper())
    
    for _ in range(4): doc.add_paragraph()
    add_centered_line(data['course'], bold=True)
    doc.add_paragraph()
    add_centered_line(f"ASSIGNMENT TOPIC:\n\"{data['title'].upper()}\"", bold=True)
    
    for _ in range(5): doc.add_paragraph()
    add_centered_line(f"BY\n{data['student_name'].upper()}", bold=True)
    add_centered_line(f"ID NUMBER: {data['student_id']}")
    
    for _ in range(2): doc.add_paragraph()
    add_centered_line(f"LECTURER: {data['lecturer'].upper()}")
    add_centered_line(f"DATE OF SUBMISSION: {data['date']}")
    
    doc.add_page_break()
    return doc

# --- Streamlit Web Interface Configuration ---
st.set_page_config(page_title="UEW Assignment Formatter", page_icon="📝")

st.title("📝 UEW Assignment Formatter")
st.write("Fill in your details below to generate a perfectly formatted Microsoft Word document according to UEW guidelines.")

# Web Form Input Fields
faculty = st.text_input("Faculty/School Name", placeholder="e.g., Faculty of Science Education")
department = st.text_input("Department Name", placeholder="e.g., Department of ICT Education")
course = st.text_input("Course Code & Title", placeholder="e.g., ICT 241: Desktop Publishing")
title = st.text_input("Assignment Topic/Title")
student_name = st.text_input("Your Full Name")
student_id = st.text_input("Student ID Number")
lecturer = st.text_input("Lecturer's Name", placeholder="e.g., Dr. John Mensah")
date = st.text_input("Submission Date", placeholder="e.g., July 20, 2026")

if st.button("Generate Assignment Template ✨"):
    if not all([faculty, department, course, title, student_name, student_id, lecturer, date]):
        st.error("Please fill in all the text fields before generating your document.")
    else:
        form_data = {
            'faculty': faculty, 'department': department, 'course': course,
            'title': title, 'student_name': student_name, 'student_id': student_id,
            'lecturer': lecturer, 'date': date
        }
        
        # Process document in memory (so users can download it via the web)
        doc = create_uew_document(form_data)
        bio = io.BytesIO()
        doc.save(bio)
        
        st.success("🎉 Document generated successfully!")
        st.download_button(
            label="💾 Download Word Document",
            data=bio.getvalue(),
            file_name=f"UEW_{course.split(':')[0].strip()}_Assignment.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )